from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "UniLink_Academic_Report_Text.md"
OUTPUT = ROOT / "UniLink_Academic_Report.docx"


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def ensure_code_style(document):
    styles = document.styles
    if "Code Block" in styles:
        return styles["Code Block"]

    style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Consolas"
    style.font.size = Pt(9)
    style.font.color.rgb = RGBColor(31, 41, 55)
    style.paragraph_format.left_indent = Inches(0.25)
    style.paragraph_format.right_indent = Inches(0.15)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    return style


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Title", 20, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        if style_name == "Heading 1":
            style.paragraph_format.space_before = Pt(16)
            style.paragraph_format.space_after = Pt(8)
        elif style_name == "Heading 2":
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        elif style_name == "Heading 3":
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(4)
        else:
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.10

    for list_style in ["List Bullet", "List Number"]:
        style = document.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    ensure_code_style(document)


def add_footer(document):
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = 1
    run = footer.add_run("UniLink Mobile Application Report")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(107, 114, 128)
    run.add_text(" | Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    footer._p.append(fld_begin)
    footer._p.append(instr)
    footer._p.append(fld_end)


def add_paragraph_with_inline_code(document, text, style=None):
    paragraph = document.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(30, 64, 175)
        else:
            paragraph.add_run(part)
    return paragraph


def build_report():
    document = Document()
    configure_document(document)
    add_footer(document)

    in_code = False
    code_lang = ""

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            in_code = not in_code
            code_lang = line[3:].strip() if in_code else ""
            if in_code and code_lang:
                p = document.add_paragraph(f"Code snippet ({code_lang})", style="Caption")
                p.runs[0].font.italic = True
            continue

        if in_code:
            p = document.add_paragraph(line if line else " ", style="Code Block")
            set_cell_shading(p, "F3F4F6")
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            document.add_paragraph(line[2:].strip(), style="Title")
            continue
        if line.startswith("## "):
            document.add_paragraph(line[3:].strip(), style="Heading 1")
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            continue
        if line.startswith("#### "):
            document.add_paragraph(line[5:].strip(), style="Heading 3")
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if number_match:
            add_paragraph_with_inline_code(
                document,
                number_match.group(2),
                style="List Number",
            )
            continue

        if line.startswith("- "):
            add_paragraph_with_inline_code(document, line[2:], style="List Bullet")
            continue

        add_paragraph_with_inline_code(document, line)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    try:
        document.save(OUTPUT)
        print(f"Built {OUTPUT}")
    except PermissionError:
        fallback = ROOT / "UniLink_Academic_Report_Updated.docx"
        document.save(fallback)
        print(f"Could not overwrite {OUTPUT}; built {fallback} instead.")


if __name__ == "__main__":
    build_report()
