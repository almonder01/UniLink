from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "UniLink_Academic_Report_Text.md"
OUTPUT = ROOT / "UniLink_Academic_Report_Final.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(107, 114, 128)
INK = RGBColor(31, 41, 55)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size=None, color=None, bold=None):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def paragraph_bottom_border(paragraph, color="D9E2EF", size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def ensure_styles(document):
    normal = document.styles["Normal"]
    set_style_font(normal, "Calibri", 11, INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_map = [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]
    for style_name, size, color, before, after in style_map:
        style = document.styles[style_name]
        set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    for list_style in ["List Bullet", "List Number"]:
        style = document.styles[list_style]
        set_style_font(style, "Calibri", 11, INK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Code Block" not in document.styles:
        code = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = document.styles["Code Block"]
    set_style_font(code, "Consolas", 9, RGBColor(17, 24, 39))
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.0


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    ensure_styles(document)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("UniLink Academic Report")
    set_run_font(run, "Calibri", 8, MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("UniLink Mobile Application Report - Final DOCX")
    set_run_font(run, "Calibri", 8, MUTED)


def add_masthead(document, title):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("ACADEMIC PROJECT REPORT")
    set_run_font(run, "Calibri", 10, BLUE, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    set_run_font(run, "Calibri", 24, NAVY, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(
        "Updated final documentation for the Flutter, Firebase, multimedia, "
        "map, search/filter, chat, notification, and admin permission workflows."
    )
    set_run_font(run, "Calibri", 11, MUTED)

    metadata = [
        ("Project", "UniLink"),
        ("Platform", "Flutter mobile application"),
        ("Backend", "Firebase Authentication and Cloud Firestore"),
        ("Updated", "July 8, 2026"),
        ("Status", "Final clean Word report"),
    ]
    for label, value in metadata:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, "Calibri", 10.5, INK, bold=True)
        r = p.add_run(value)
        set_run_font(r, "Calibri", 10.5, INK)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(14)
    paragraph_bottom_border(rule)


def add_paragraph_with_inline_code(document, text, style=None):
    paragraph = document.add_paragraph(style=style)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 9.5, RGBColor(30, 64, 175))
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Calibri", 11, INK)
    return paragraph


def add_code_line(document, line):
    paragraph = document.add_paragraph(line if line else " ", style="Code Block")
    paragraph_shading(paragraph, "F3F4F6")


def build_report():
    document = Document()
    configure_document(document)

    in_code = False
    title_written = False

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            in_code = not in_code
            language = line[3:].strip()
            if in_code and language:
                caption = document.add_paragraph()
                caption.paragraph_format.space_before = Pt(4)
                caption.paragraph_format.space_after = Pt(2)
                run = caption.add_run(f"Code snippet ({language})")
                set_run_font(run, "Calibri", 9, MUTED, italic=True)
            continue

        if in_code:
            add_code_line(document, line)
            continue

        if not line.strip():
            continue

        if line.startswith("# "):
            add_masthead(document, line[2:].strip())
            title_written = True
            continue
        if line.startswith("## "):
            if not title_written:
                add_masthead(document, "UniLink Mobile Application Report")
                title_written = True
            document.add_paragraph(line[3:].strip(), style="Heading 1")
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:].strip(), style="Heading 2")
            continue
        if line.startswith("#### "):
            document.add_paragraph(line[5:].strip(), style="Heading 3")
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", line)
        if number_match:
            add_paragraph_with_inline_code(
                document,
                number_match.group(2),
                style="List Number",
            )
            continue

        if line.startswith("- "):
            add_paragraph_with_inline_code(document, line[2:], style="List Bullet")
            continue

        add_paragraph_with_inline_code(document, line)

    document.save(OUTPUT)
    print(f"Built {OUTPUT}")


if __name__ == "__main__":
    build_report()
